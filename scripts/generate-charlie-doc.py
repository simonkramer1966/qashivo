#!/usr/bin/env python3
"""Generate CHARLIE_TECHNICAL_ARCHITECTURE.docx — comprehensive technical architecture document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── Style setup ──────────────────────────────────────────────────────────────

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.styles['Heading 1'].font.size = Pt(20)
doc.styles['Heading 2'].font.size = Pt(15)
doc.styles['Heading 3'].font.size = Pt(12)

# Create a code style
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(4)
code_style.paragraph_format.line_spacing = 1.0

def add_code(text):
    """Add a code block."""
    for line in text.strip().split('\n'):
        p = doc.add_paragraph(line, style='CodeBlock')
    doc.add_paragraph()  # spacing after

def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(9)
    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

def add_bold_para(bold_text, normal_text):
    """Paragraph starting with bold text followed by normal text."""
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    p.add_run(normal_text)


# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CHARLIE')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Autonomous Collections Decision Engine')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x6a)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Technical Architecture Document\n').font.size = Pt(14)
meta.add_run('Qashivo Ltd\n\n').font.size = Pt(12)
meta.add_run('Version 1.1 — April 2026\n').font.size = Pt(11)
meta.add_run('Classification: Internal / Technical Leadership\n\n').font.size = Pt(10)
meta.add_run('Audience: CTO, Principal Engineers, Technical Due Diligence').font.size = Pt(10)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Table of Contents', level=1)
toc_entries = [
    '1. Executive Summary',
    '2. System Architecture Overview',
    '3. Decision Engine',
    '    3.1 Priority Scoring',
    '    3.2 Channel Selection',
    '    3.3 Tone Escalation',
    '    3.4 Timing & Composite Scoring',
    '    3.5 Invoice Consolidation',
    '4. Learning & Intelligence Systems',
    '    4.1 Promise Reliability Score (PRS)',
    '    4.2 Channel Effectiveness',
    '    4.3 Payment Attribution',
    '    4.4 Portfolio Urgency',
    '    4.5 P(Pay) Distribution Model',
    '    4.6 Debtor Enrichment',
    '    4.7 Cold Start & Segment Priors',
    '    4.8 Seasonal Patterns',
    '    4.9 Debtor Grouping',
    '5. Message Generation Pipeline',
    '    5.1 LLM Generation',
    '    5.2 Output Validation',
    '    5.3 Circuit Breaker',
    '    5.4 Template Fallback',
    '    5.5 Compliance Engine',
    '6. Execution & Delivery Pipeline',
    '    6.1 Scheduler Architecture',
    '    6.2 Execution-Time Validation Gate',
    '    6.3 Delivery Tracking',
    '    6.4 Retry Logic',
    '    6.5 Legal Evidence & Voice Contact Records',
    '7. Compliance & Safety Systems',
    '8. Data Architecture',
    '9. Integration Points',
    '10. What Makes This System Exceptional',
]
for entry in toc_entries:
    p = doc.add_paragraph(entry)
    p.paragraph_format.space_after = Pt(2)
    if not entry.startswith('    '):
        p.runs[0].bold = True

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'Charlie is Qashivo\'s autonomous collections decision engine. It is not a template mailer, not a scheduled reminder system, '
    'and not a rules engine with an AI label bolted on. Charlie is a closed-loop learning system that decides when to contact each '
    'debtor, through which channel, in what tone, and with what message \u2014 then learns from the outcome to make a better decision next time.'
)

doc.add_paragraph(
    'The system operates a continuous learning cycle. Overdue invoices enter Charlie\'s priority scoring pipeline, which evaluates '
    'every invoice against a composite formula weighing payment probability, communication friction, compliance risk, and portfolio '
    'urgency. The highest-scoring candidates receive actions planned by the adaptive scheduler, which fits a log-normal distribution '
    'to each debtor\'s historical payment behaviour and evaluates six time horizons to find the optimal moment to make contact. '
    'Messages are generated by Claude (Anthropic\'s LLM) with full business context \u2014 debtor history, invoice details, prior '
    'conversations, tone level, and seasonal patterns \u2014 producing genuinely personalised communications that pass through a '
    'five-point output validator and six-rule compliance engine before delivery.'
)

doc.add_paragraph(
    'After delivery, the feedback loop closes. SendGrid webhooks report delivery status, opens, clicks, and replies. Xero sync '
    'detects invoice payment transitions. These signals flow through a three-tier channel effectiveness model (delivery 0.2, '
    'engagement 0.2, payment 0.6) that updates each debtor\'s learning profile via adaptive exponential moving averages. The '
    'Promise Reliability Score tracks whether debtors keep their promises using Bayesian inference with recency weighting. '
    'Portfolio-level urgency recomputes nightly, adjusting pressure based on DSO trajectory. Every decision Charlie makes '
    'tomorrow is informed by what it learned today.'
)

doc.add_paragraph(
    'What differentiates Charlie from every collections automation product on the market is this: the system has no templates. '
    'Every message is generated fresh by an LLM with full context. Every timing decision is backed by a statistical model fitted to '
    'that specific debtor. Every tone escalation is velocity-capped and compliance-checked. Every channel selection reflects learned '
    'effectiveness, not configuration. And the entire pipeline is legally defensible \u2014 every action is logged with reasoning, '
    'compliance checks are auditable, and a 30-day statutory response window is enforced before any legal escalation.'
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('2. System Architecture Overview', level=1)

doc.add_paragraph(
    'The end-to-end data flow from "invoice becomes overdue" to "debtor pays" traverses seven distinct layers, each with a '
    'specific responsibility. Understanding this flow is essential to understanding Charlie.'
)

doc.add_heading('The Seven Layers', level=2)

doc.add_paragraph(
    'Layer 1: Data Ingestion. Xero sync (server/services/xeroSync.ts) runs on a 4-hour background cadence, fetching invoices '
    'and contacts via the Xero API with 1.5-second rate limiting between pages. Invoices are written to both cached_xero_invoices '
    '(raw reference) and the main invoices table. Contacts are upserted, but AR overlay fields (arContactEmail, arContactPhone, '
    'arContactName, arNotes) are never overwritten \u2014 these are Qashivo-owned data. When an invoice transitions to PAID status '
    'during sync, the system fires payment attribution through the event bus, closing the feedback loop.'
)

doc.add_paragraph(
    'Layer 2: Intelligence & Scoring. The portfolio controller (server/services/portfolioController.ts) orchestrates nightly '
    'processing at 2 AM. It calls recomputeUrgency() in the Charlie decision engine to adjust tenant-level urgency based on DSO '
    'trajectory, then calculatePerDebtorUrgency() to weight individual debtors by their contribution to overdue and payment trend. '
    'The Promise Reliability Service recalculates PRS scores with Bayesian adjustment. Debtor enrichment runs quarterly via '
    'Companies House API integration.'
)

doc.add_paragraph(
    'Layer 3: Decision Engine. The Charlie decision engine (server/services/charlieDecisionEngine.ts) evaluates each overdue '
    'invoice through priority scoring, channel selection, tone escalation, and timing optimisation. The output is a CharlieDecision '
    'object containing the recommended action: channel, tone, timing, confidence, and whether human review is required.'
)

doc.add_paragraph(
    'Layer 4: Action Planning. The action planner (server/services/actionPlanner.ts) converts decisions into concrete actions. '
    'It consolidates multiple invoices per debtor into single communications, applies debtor group consistency (Gap 12), enforces '
    'channel preference overrides (Gap 11), and integrates seasonal adjustments (Gap 13). Actions enter the system as either '
    'pending_approval (semi-auto mode) or scheduled (full-auto mode).'
)

doc.add_paragraph(
    'Layer 5: Message Generation. The AI message generator (server/services/aiMessageGenerator.ts) produces email, SMS, and '
    'voice content via Claude API calls. A three-state circuit breaker protects against LLM outages. Generated content passes '
    'through a five-point output validator checking length, debtor name, invoice references, system prompt leakage, and tone '
    'alignment. If the circuit breaker is open beyond 4 hours, the system falls back to 10 static templates (5 tones x 2 channels).'
)

doc.add_paragraph(
    'Layer 6: Compliance & Execution. The compliance engine (server/services/compliance/complianceEngine.ts) applies six rules '
    'before any message reaches the wire: frequency cap, channel cooldown, time-of-day, prohibited language, data isolation, and '
    'debtor vulnerability. The action executor (server/services/actionExecutor.ts) runs a validation gate at execution time '
    '\u2014 re-checking invoice status, legal windows, and probable payments \u2014 before routing to the appropriate delivery wrapper '
    '(SendGrid for email, Vonage for SMS, Retell AI for voice). All wrappers enforce communication mode (Off/Testing/Soft Live/Live) '
    'and fail closed on errors.'
)

doc.add_paragraph(
    'Layer 7: Feedback & Learning. SendGrid delivery webhooks flow through the event bus (server/services/event-bus.ts) to update '
    'action delivery status and trigger channel effectiveness recalculation. Payment detection during Xero sync triggers attribution '
    'analysis. The channel effectiveness service updates learning profiles via adaptive EMA. This closes the loop: tomorrow\'s '
    'decisions reflect today\'s outcomes.'
)

doc.add_heading('Service Map', level=2)

add_table(
    ['Service', 'File', 'Role'],
    [
        ['Charlie Decision Engine', 'server/services/charlieDecisionEngine.ts', 'Priority scoring, channel selection, segment determination, portfolio urgency'],
        ['Tone Escalation Engine', 'server/services/toneEscalationEngine.ts', 'Five-level tone with velocity cap, no-response escalation, vulnerable ceiling'],
        ['Action Planner', 'server/services/actionPlanner.ts', 'Action creation, consolidation, debtor group enforcement, AI optimisation'],
        ['Action Executor', 'server/services/actionExecutor.ts', 'Execution-time validation, delivery, retry logic, legal window'],
        ['Collections Scheduler', 'server/services/collectionsScheduler.ts', 'Two-phase orchestrator: hourly planner, 10-minute executor'],
        ['Portfolio Controller', 'server/services/portfolioController.ts', 'Nightly urgency, 6-hourly adaptive planning, daily plan'],
        ['Adaptive Scheduler', 'server/lib/adaptive-scheduler.ts', 'Composite scoring: P(Pay), friction, compliance risk, urgency'],
        ['AI Message Generator', 'server/services/aiMessageGenerator.ts', 'LLM generation with circuit breaker + validation + fallback'],
        ['LLM Circuit Breaker', 'server/services/llmCircuitBreaker.ts', 'Three-state machine: closed/open/half_open'],
        ['LLM Output Validator', 'server/services/llmOutputValidator.ts', 'Five-point quality check on generated content'],
        ['Template Fallback', 'server/services/templateFallback.ts', 'Ten static templates (5 tones x 2 channels) for circuit-open fallback'],
        ['Compliance Engine', 'server/services/compliance/complianceEngine.ts', 'Six-rule compliance gate with audit logging'],
        ['Promise Reliability', 'server/services/promiseReliabilityService.ts', 'Bayesian PRS with recency weighting'],
        ['Channel Effectiveness', 'server/services/channelEffectivenessService.ts', 'Three-tier model, adaptive EMA, payment attribution'],
        ['Payment Distribution', 'server/services/paymentDistribution.ts', 'Log-normal distribution, seasonal adjustments, forecasting'],
        ['Collection Learning', 'server/services/collectionLearningService.ts', 'Segment priors, cold start, A/B testing, profile management'],
        ['Debtor Enrichment', 'server/services/debtorEnrichmentService.ts', 'Companies House integration, credit risk scoring'],
        ['Legal Window Job', 'server/jobs/legalWindowJob.ts', '30-day statutory window monitoring, expiry warnings'],
    ]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 3. DECISION ENGINE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('3. Decision Engine', level=1)

doc.add_paragraph(
    'The decision engine is the core of Charlie. It answers four questions for every overdue invoice: how urgent is this '
    '(priority scoring), how should we reach the debtor (channel selection), what posture should the communication take '
    '(tone escalation), and when is the optimal moment to make contact (timing). The answers are synthesised into a '
    'CharlieDecision object that drives the entire downstream pipeline.'
)

# ── 3.1 Priority Scoring ─────────────────────────────────────────────────────

doc.add_heading('3.1 Priority Scoring', level=2)

doc.add_paragraph(
    'Priority scoring in charlieDecisionEngine.ts assigns each invoice a numeric score (0\u2013100) and a tier '
    '(critical, high, medium, low, or excluded). The algorithm runs in five phases.'
)

add_bold_para('Phase 1: Exclusion. ', 'Invoices in admin_blocked, disputed, or ptp (promise-to-pay active) states '
    'are immediately assigned to the excluded tier with a score of 0. Invoices on hold or in terminal states '
    '(ptp_met, debt_recovery) are also excluded. These never generate actions.')

add_bold_para('Phase 2: Critical fast-track. ', 'A missed promise-to-pay is the highest-priority event in the system. '
    'Any invoice where the debtor made a promise and broke it receives a score of 95 and the critical tier. This '
    'bypasses all other scoring logic.')

add_bold_para('Phase 3: High priority. ', 'Two conditions qualify for high priority: invoices 60+ days overdue with '
    'a value of \u00a310,000 or more score 85; invoices 30\u201360 days overdue with a value of \u00a35,000 or more score 75.')

add_bold_para('Phase 4: Medium priority. ', 'Newly overdue invoices (0\u20137 days) that are either high value or '
    'belong to new customers score 65. Early intervention on significant amounts prevents them from ageing.')

add_bold_para('Phase 5: Accumulated scoring. ', 'All remaining invoices are scored by accumulating points across '
    'four dimensions:')

add_table(
    ['Component', 'Condition', 'Points'],
    [
        ['Overdue severity', '60+ days', '+40'],
        ['', '30\u201360 days', '+30'],
        ['', '14\u201330 days', '+20'],
        ['', '7\u201314 days', '+10'],
        ['', '0\u20137 days', '+5'],
        ['Amount', '\u2265 \u00a310,000', '+15'],
        ['', '\u2265 \u00a35,000', '+10'],
        ['Contact recency', '14+ days since last contact (overdue)', '+15'],
        ['', 'Never contacted (overdue)', '+10'],
        ['', '7+ days since last contact', '+5'],
        ['State escalation', 'final_demand state', '+20'],
    ]
)

doc.add_paragraph(
    'The accumulated score maps to tiers: 80+ is high, 50\u201379 is medium, below 50 is low. The tier determines '
    'how aggressively the action planner schedules follow-up.'
)

# ── 3.2 Channel Selection ────────────────────────────────────────────────────

doc.add_heading('3.2 Channel Selection', level=2)

doc.add_paragraph(
    'Channel selection follows a mandated progression: email first, then SMS, then voice. This progression is not '
    'arbitrary \u2014 email creates a written audit trail (essential for legal escalation), SMS has higher open rates '
    'for short notices, and voice is reserved for high-value or persistently unresponsive debtors where a direct '
    'conversation may break a deadlock.'
)

doc.add_paragraph(
    'The algorithm tracks prior attempts through the invoice\'s reminderCount field. First touch (reminderCount 0 or '
    'no prior contact date) always uses email if an email address is available, falling back to SMS if only a phone number '
    'exists. After 48\u201372 hours with no response and at least one prior email, the system escalates to SMS. After 5+ '
    'days with no response and at least one prior SMS, voice becomes available \u2014 but only for invoices that are either '
    'high value (\u2265\u00a310,000) or 30+ days overdue. Enterprise-segment debtors always prefer email throughout the cycle '
    '(formal communication norms). Final demand actions always use email (written notice requirement).'
)

doc.add_paragraph(
    'Channel preference hard overrides (Gap 11) are applied after the algorithmic selection. The system reads '
    'emailEnabled, smsEnabled, and voiceEnabled flags from the customerPreferences table. A null or undefined value '
    'defaults to true (backwards compatible). If the selected channel is disabled, the system falls back in order: '
    'email, then SMS, then voice. If all channels are disabled, the action is skipped entirely. These overrides can '
    'be set manually by users or extracted automatically by Riley from debtor conversations ("don\'t call me", '
    '"email only"). The applyChannelPreferenceOverride() method in charlieDecisionEngine.ts handles this logic.'
)

# ── 3.3 Tone Escalation ──────────────────────────────────────────────────────

doc.add_heading('3.3 Tone Escalation', level=2)

doc.add_paragraph(
    'The tone escalation engine (server/services/toneEscalationEngine.ts) manages a five-level tone scale: friendly, '
    'professional, firm, formal, and legal. The determineTone() function takes a ToneEscalationInput (tenantId, contactId, '
    'daysOverdue, touchCount) and returns a ToneEscalationResult containing the resolved tone level, reasoning chain, '
    'and detailed signals.'
)

doc.add_heading('Baseline Determination', level=3)

doc.add_paragraph(
    'The baseline tone is determined by mapping days overdue against escalation thresholds, which vary by tenant style '
    '(GENTLE, STANDARD, or FIRM). Under the STANDARD profile, invoices stay at friendly for the first 7 days overdue, '
    'professional from 7\u201314 days, firm from 14\u201330 days, formal from 30\u201360 days, and legal beyond 60 days. '
    'The GENTLE profile extends each window (friendly to 14 days, professional to 21), while FIRM compresses them '
    '(friendly to 4 days, professional to 11).'
)

add_table(
    ['Tone Level', 'GENTLE (days)', 'STANDARD (days)', 'FIRM (days)'],
    [
        ['Friendly', '0\u201314', '0\u20137', '0\u20134'],
        ['Professional', '14\u201321', '7\u201314', '4\u201311'],
        ['Firm', '21\u201337', '14\u201330', '11\u201327'],
        ['Formal', '37\u201367', '30\u201360', '27\u201357'],
        ['Legal', '67+', '60+', '57+'],
    ]
)

doc.add_heading('Signal-Based Adjustments', level=3)

doc.add_paragraph(
    'Six adjustments modify the baseline, each capable of stepping the tone up or down by one level. '
    'A recently engaged debtor (inbound message within 7 days, per RECENT_ENGAGEMENT_DAYS) steps down one level \u2014 '
    'there is no point escalating when the debtor is actively communicating. A serial promiser (flagged by the PRS '
    'service when promiseSequence \u2265 3 and PRS < 60) steps up one level. A previously good payer (PRS \u2265 70, '
    'per GOOD_PAYER_PRS_THRESHOLD) with low touch count (\u22642) who is not a serial promiser is held at friendly \u2014 '
    'these debtors deserve patience. New customers (created < 90 days ago) with one or fewer touches are capped at professional. '
    'A deteriorating relationship (PRS trend shows 15+ point drop over 90 days) steps up one level. Finally, no-response '
    'escalation pressure (Gap 5) steps up one level when the debtor has ignored 4+ consecutive contacts '
    '(DEFAULT_NO_RESPONSE_THRESHOLD = 4), counted by getConsecutiveNoResponseCount() which checks the last 10 completed actions '
    'for inbound email responses.'
)

doc.add_heading('Velocity Cap', level=3)

doc.add_paragraph(
    'The velocity cap (Gap 5) prevents tone from jumping more than one level between contact cycles. The engine reads '
    'agentToneLevel from the most recent completed action via getLastSentAction() (filtering out failed and bounced '
    'deliveries). The calculated tone cannot exceed lastTone + 1 or fall below lastTone \u2212 1. On first contact, '
    'no velocity cap is applied.'
)

doc.add_paragraph(
    'One exception bypasses the downward velocity cap: the significant payment override. If the debtor has paid 50% '
    'or more of their outstanding balance since the last action (DEFAULT_SIGNIFICANT_PAYMENT_THRESHOLD = 0.50), the '
    'baseline resets to professional regardless of where tone was previously. This prevents a debtor who has just made '
    'a substantial payment from receiving a firm or formal follow-up on the remaining balance.'
)

doc.add_heading('Hard Caps', level=3)

doc.add_paragraph(
    'Two hard caps override all other logic. Vulnerable debtors (contacts with isPotentiallyVulnerable = true) are '
    'capped at professional (VULNERABLE_CEILING). GENTLE-style tenants cannot use the legal tone \u2014 it is capped '
    'to formal. The ToneEscalationResult includes wasCapped and uncappedTone fields for audit transparency.'
)

# ── 3.4 Timing & Composite Scoring ───────────────────────────────────────────

doc.add_heading('3.4 Timing & Composite Scoring', level=2)

doc.add_paragraph(
    'The adaptive scheduler (server/lib/adaptive-scheduler.ts) determines when to contact each debtor by evaluating '
    'six time horizons (12h, 24h, 36h, 48h, 72h, 96h from now) across all allowed channels. For each (time, channel) '
    'candidate, it computes a composite score using four weighted components.'
)

doc.add_paragraph(
    'The composite scoring formula is:'
)

add_code('Score = 1.0 * P(Pay) - 0.35 * Friction - 0.6 * Risk + (0.4 + urgencyFactor) * Urgency')

doc.add_paragraph(
    'The weights are fixed: alpha = 1.0 for payment probability, beta = 0.35 for friction cost, gamma = 0.6 for '
    'compliance risk, and delta = 0.4 + urgencyFactor for urgency (where urgencyFactor is the per-debtor urgency '
    'from the portfolio controller, typically 0.1\u20131.0). The candidate with the highest composite score wins.'
)

doc.add_heading('P(Pay) Estimation', level=3)

doc.add_paragraph(
    'P(Pay) is the probability that the debtor will pay within the given time horizon, conditional on not having paid yet. '
    'The estimateProbabilityOfPayment() function in adaptive-scheduler.ts first fits a log-normal distribution to the debtor\'s '
    'historical payment data via fitDistribution() in paymentDistribution.ts. The distribution\'s mu (location) parameter comes '
    'from the trend-adjusted median days-to-pay, and sigma (spread) comes from the p75/median ratio or volatility fallback. '
    'Seasonal adjustments shift mu for the current month. The base probability uses the conditional CDF: '
    'P(pay in horizon | not paid yet) = (CDF(currentDay + horizon) \u2212 CDF(currentDay)) / (1 \u2212 CDF(currentDay)).'
)

doc.add_paragraph(
    'Five additive modifiers refine the base probability. Channel effectiveness boost adds up to 60% of the channel\'s '
    'historical reply rate. An amount sensitivity penalty (log10(amount)/20, capped at 0.25) reduces probability for '
    'larger invoices. A weekday effect uses learned day-of-week payment patterns from customerBehaviorSignals. '
    'A promise-to-pay boost of 0.15 applies when the contact time falls before the promised payment date. '
    'The final probability is clamped to [0, 1].'
)

doc.add_heading('Friction Cost', level=3)

doc.add_paragraph(
    'Friction measures how annoying this contact will be to the debtor. The base friction is 0.7 minus the channel\'s '
    'reply rate (new customers default to 0.7 \u2014 maximum friction). A repeat penalty of 0.1 applies if the proposed '
    'channel is the same as the last channel used. High friction reduces the composite score, steering the system '
    'toward channels the debtor has historically engaged with.'
)

doc.add_heading('Compliance Risk', level=3)

doc.add_paragraph(
    'Compliance risk scores the legal and reputational danger of contacting at this time. Disputed invoices carry 0.8 risk '
    '(nearly blocking). Calling before an invoice is due scores 0.15. Contacting during an active promise-to-pay window '
    'scores 0.3. The baseline risk for standard contacts is 0.05.'
)

doc.add_heading('Urgency Boost', level=3)

doc.add_paragraph(
    'Urgency applies portfolio-level pressure. The formula urgencyFactor / sqrt(hoursUntilTouch) favours earlier contact '
    'times when urgency is high. The delta weight (0.4 + urgencyFactor) means that a debtor with urgencyFactor 0.8 '
    'gets an urgency weight of 1.2, while a low-priority debtor at 0.2 gets only 0.6. This ensures high-urgency debtors '
    'are contacted sooner without overwhelming the system.'
)

doc.add_heading('Guard Conditions', level=3)

doc.add_paragraph(
    'The scheduler returns null (no scheduling) if: a manual override exists, the invoice is disputed, the daily touch '
    'limit for the workflow or customer has been reached, or no channels are available after preference filtering. '
    'All candidate times must pass quiet hours validation (respecting tenant timezone) and minimum gap checks '
    '(configurable hours between touches for the same debtor).'
)

# ── 3.5 Invoice Consolidation ────────────────────────────────────────────────

doc.add_heading('3.5 Invoice Consolidation', level=2)

doc.add_paragraph(
    'The action planner\'s planAdaptiveActions() function consolidates multiple invoices per contact into single '
    'communications. After scoring all invoices individually, it groups them by contact and creates one bundled action '
    'per contact using the highest-priority recommendation. Only invoices that pass all exclusion checks (not disputed, '
    'not on hold, not paid, not paused) enter the consolidated action. If the consolidated total falls below the tenant\'s '
    'minimum chase threshold, the action is skipped entirely. This prevents sending a collections email for a trivial '
    'outstanding balance.'
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 4. LEARNING & INTELLIGENCE SYSTEMS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('4. Learning & Intelligence Systems', level=1)

doc.add_paragraph(
    'Charlie\'s learning systems are what elevate it from a decision engine to an adaptive intelligence. Nine subsystems '
    'work together to build an increasingly accurate model of each debtor\'s behaviour, payment patterns, and '
    'communication preferences.'
)

# ── 4.1 PRS ──────────────────────────────────────────────────────────────────

doc.add_heading('4.1 Promise Reliability Score (PRS)', level=2)

doc.add_paragraph(
    'The Promise Reliability Service (server/services/promiseReliabilityService.ts) tracks whether debtors keep their '
    'promises. When a debtor makes a promise to pay (via agent conversation, Riley extraction, or manual entry), '
    'createPromise() records it. When the promise is evaluated (kept, broken, partially_kept, or cancelled), '
    'evaluatePromise() updates the record and triggers PRS recalculation.'
)

doc.add_paragraph(
    'The PRS calculation uses Bayesian inference with recency weighting. Each resolved promise is scored (kept = 100, '
    'partially_kept = 50, broken = 0) and weighted by a 90-day half-life decay: weight = 1 / (1 + daysSinceResolution / 90). '
    'A promise resolved today has weight 1.0; a promise resolved 90 days ago has weight 0.5; a promise from 6 months ago '
    'has weight 0.25. The raw PRS is the weighted average: sum(weight * score) / sum(weight).'
)

doc.add_paragraph(
    'Bayesian adjustment prevents thin-data debtors from receiving extreme scores. With BAYESIAN_K = 3 (equivalent '
    'to three promises of evidence), the adjusted PRS is (totalWeight * rawPRS + K * prior) / (totalWeight + K). '
    'The prior is the tenant population mean PRS if the tenant has 10+ scored debtors (MIN_TENANT_DEBTORS_FOR_MEAN), '
    'otherwise the system default of 60 (SYSTEM_DEFAULT_PRIOR). This means a debtor who has kept one promise scores '
    'approximately 70 (not 100), while a debtor with zero history starts at 60. Confidence is calculated as '
    'totalWeight / (totalWeight + K), approaching 1.0 as evidence accumulates.'
)

doc.add_paragraph(
    'The service also derives four behavioural flags from the PRS data. isSerialPromiser triggers when a debtor has '
    'a promise sequence of 3+ and adjusted PRS below 60. isReliableLatePayer flags debtors with PRS between 40 and 70 '
    'who have payment history (they pay, just slowly). isRelationshipDeteriorating fires when the 30-day PRS is 15+ '
    'points below the 90-day PRS. isNewCustomer flags debtors with fewer than 3 total promises. These flags feed into '
    'tone escalation and channel selection decisions.'
)

# ── 4.2 Channel Effectiveness ────────────────────────────────────────────────

doc.add_heading('4.2 Channel Effectiveness', level=2)

doc.add_paragraph(
    'The channel effectiveness service (server/services/channelEffectivenessService.ts) maintains a three-tier model '
    'that scores how effective each channel (email, SMS, voice) is for each debtor. The three tiers, weighted by importance, are:'
)

add_table(
    ['Tier', 'Weight', 'Signals', 'Score Range'],
    [
        ['Delivery', '0.2', 'delivered=1.0, bounced=0.0, unknown=0.5', '0\u20131.0'],
        ['Engagement', '0.2', 'replied=0.9, clicked=0.5, opened=0.3 (max, not sum)', '0\u20131.0'],
        ['Payment', '0.6', 'Attribution score from payment analysis', '0\u20131.0'],
    ]
)

doc.add_paragraph(
    'The combined effectiveness = (delivery * 0.2) + (engagement * 0.2) + (payment * 0.6). Payment outcome '
    'dominates because the ultimate measure of a collection channel\'s effectiveness is whether the debtor pays.'
)

doc.add_paragraph(
    'Scores are updated via adaptive exponential moving average (EMA). The EMA retention factor is tied to confidence: '
    'retention = 0.5 + (0.3 * confidence), ranging from 0.53 at low confidence to 0.77 at high confidence. This means '
    'new debtors\' scores change quickly (low retention, high responsiveness to new data), while established profiles '
    'change gradually (high retention, stable scores). Confidence itself grows by 0.1 * (1 - confidence) per update, '
    'capped at 0.95.'
)

doc.add_paragraph(
    'Hard bounces bypass the EMA entirely. When a hard bounce is detected (via SendGrid webhook), '
    'handleHardBounceEffectiveness() immediately sets email effectiveness to 0.1. This is a definitive signal that '
    'the email address is invalid \u2014 no amount of averaging should dilute it.'
)

# ── 4.3 Payment Attribution ──────────────────────────────────────────────────

doc.add_heading('4.3 Payment Attribution', level=2)

doc.add_paragraph(
    'When Xero sync detects an invoice has transitioned to PAID, the system must attribute that payment to the '
    'collection actions that influenced it. calculatePaymentAttribution() in channelEffectivenessService.ts implements '
    'a time-windowed attribution model with three configurable parameters (stored as tenant settings): '
    'paymentAttributionSameDayExcluded (default true), paymentAttributionFullCreditHours (default 48), and '
    'paymentAttributionPartialCreditDays (default 7).'
)

doc.add_paragraph(
    'Same-day exclusion recognises that if a debtor pays the same day they receive a collection message, the payment '
    'was likely already in flight \u2014 the message did not cause it. Within the full credit window (48 hours), '
    'the most recent action receives full attribution (1.0). Between 48 hours and 7 days, the most recent action '
    'receives partial credit (0.5). Beyond 7 days, no attribution is assigned.'
)

doc.add_paragraph(
    'Multi-channel attribution handles the common case where a debtor receives an email, then an SMS, then pays. '
    'processMultiChannelAttribution() gives the primary (most recent) action full or partial credit as above, and '
    'awards 0.3 credit to earlier actions within the same attribution window. This recognises that the email may '
    'have planted the seed, even though the SMS was the final nudge.'
)

# ── 4.4 Portfolio Urgency ────────────────────────────────────────────────────

doc.add_heading('4.4 Portfolio Urgency', level=2)

doc.add_paragraph(
    'The portfolio controller manages urgency at two levels: tenant-wide and per-debtor.'
)

doc.add_paragraph(
    'Tenant-wide urgency is recomputed nightly by recomputeUrgency() in charlieDecisionEngine.ts. It compares the '
    'tenant\'s projected DSO (from server/lib/dso.ts) against the target DSO (configurable, default 45 days). '
    'If projected DSO exceeds the target by more than 1 day, urgency increases by 0.1 (up to a maximum of 1.0). '
    'If projected DSO is below target by more than 1 day, urgency decreases by 0.1 (down to a minimum of 0.1). '
    'This creates a self-regulating feedback loop: when collections are effective, pressure eases; when DSO creeps up, '
    'pressure increases.'
)

doc.add_paragraph(
    'Per-debtor urgency (Gap 3) adds granularity. calculatePerDebtorUrgency() computes each debtor\'s contribution '
    'to the total overdue balance, normalised so the average contribution weight is 1.0. A debtor responsible for 20% '
    'of overdue in a portfolio of 10 debtors has a contribution weight of 2.0. This is multiplied by a trend multiplier: '
    'debtors with deteriorating payment trends (trend > 0) get multiplied by 1.0 to 1.5, while improving debtors (trend < 0) '
    'get multiplied by 0.5 to 1.0. The final per-debtor urgency = baseUrgency * contributionWeight * trendMultiplier, '
    'written to customerLearningProfiles for the action planner to read.'
)

doc.add_paragraph(
    'The effect: a small-balance debtor with improving payment trends gets near-zero urgency and may not be contacted at all, '
    'while a large-balance debtor dragging DSO upward receives maximum pressure. This prevents the common collections '
    'antipattern of treating all debtors equally regardless of their impact on the portfolio.'
)

# ── 4.5 P(Pay) Distribution Model ────────────────────────────────────────────

doc.add_heading('4.5 P(Pay) Distribution Model', level=2)

doc.add_paragraph(
    'Charlie models each debtor\'s payment behaviour as a log-normal distribution (server/services/paymentDistribution.ts). '
    'The log-normal is the correct distribution for payment timing: it is always positive (you cannot pay in negative days), '
    'right-skewed (most payments cluster near the median with a long tail of late payments), and its parameters have '
    'intuitive interpretations.'
)

doc.add_paragraph(
    'The fitDistribution() function derives distribution parameters from four signals in customerBehaviorSignals. '
    'Mu (location) is ln(adjustedMedian) where adjustedMedian = max(1, medianDaysToPay + trend * TREND_WEIGHT). '
    'TREND_WEIGHT = 3 means each unit of trend shifts the expected payment date by approximately 3 days. '
    'Sigma (spread) is derived from p75DaysToPay when available: sigma = ln(p75/median) / 0.6745. If p75 is unavailable, '
    'volatility is used as a fallback: sigma = volatility * 0.3. Sigma is clamped to [0.1, 1.5] to prevent degenerate '
    'distributions.'
)

doc.add_paragraph(
    'Three-scenario forecasting via getPaymentForecast() extracts the 25th, 50th, and 75th percentiles as optimistic, '
    'expected, and pessimistic payment dates. These feed directly into the Weekly CFO Review\'s cashflow projections, '
    'replacing the previous hardcoded multiplier approach (which used due date for optimistic, average for expected, '
    'and 1.5x average for pessimistic).'
)

# ── 4.6 Debtor Enrichment ────────────────────────────────────────────────────

doc.add_heading('4.6 Debtor Enrichment', level=2)

doc.add_paragraph(
    'The debtor enrichment service (server/services/debtorEnrichmentService.ts) augments debtor profiles with external '
    'intelligence from Companies House. When new contacts are synced from Xero, enrichNewContacts() fires non-blocking '
    'enrichment for each contact. The pipeline: search Companies House by company name, fetch the company profile and '
    'filing history, calculate a credit risk score, generate a Claude AI risk summary, and store everything in the '
    'debtorIntelligence table.'
)

doc.add_paragraph(
    'Credit risk scoring starts at a baseline of 50 and adjusts based on signals. Positive signals: company age > 10 years '
    '(+15), age > 5 years (+10), zero late filings (+10), medium or large company size (+5). Negative signals: '
    'late filing count (minus 3 per filing, capped at \u221215), insolvency risk (\u221250), dissolved status (\u221250), '
    'dormant status (\u221220), company age < 2 years (\u221210). The final score is clamped to [0, 100] and written to '
    'both debtorIntelligence.creditRiskScore and contacts.riskScore for Charlie to read.'
)

doc.add_paragraph(
    'A 90-day freshness check prevents re-enrichment of recently scored debtors. Quarterly re-enrichment '
    '(runQuarterlyEnrichment) processes up to 50 debtors per run, targeting those with outstanding balances whose '
    'enrichment is missing or stale. The Companies House API client (server/services/companiesHouse.ts) uses Basic '
    'authentication, respects the 600 requests per 5 minutes rate limit, and handles 429 responses with a 60-second '
    'wait and single retry.'
)

# ── 4.7 Cold Start & Segment Priors ──────────────────────────────────────────

doc.add_heading('4.7 Cold Start & Segment Priors', level=2)

doc.add_paragraph(
    'New debtors present a cold-start problem: with no historical data, how does Charlie make informed decisions? '
    'The collection learning service (server/services/collectionLearningService.ts) solves this through a four-level '
    'fallback chain of segment priors.'
)

doc.add_paragraph(
    'When getOrCreateCustomerProfile() creates a new learning profile, it calls getSegmentPriors() to inherit channel '
    'effectiveness scores from similar debtors. The fallback chain tries: (1) same segment and same size band within the '
    'tenant (confidence 0.20), (2) same segment across all sizes (confidence 0.15), (3) same size band across all segments '
    '(confidence 0.15), (4) system defaults of email=0.6, sms=0.5, voice=0.4 (confidence 0.10). Each level requires at '
    'least 5 mature profiles (learningConfidence > 0.5) to form a cohort. Size bands are determined by credit limit: '
    'small (<\u00a310k), medium (\u00a310k\u2013\u00a350k), large (\u2265\u00a350k).'
)

doc.add_paragraph(
    'For payment timing, the P(Pay) distribution model uses segment-specific log-normal priors: small_business has '
    'mu = ln(40) with sigma = 0.5 (40-day typical payment), enterprise has mu = ln(50) with sigma = 0.4 (slower but '
    'more consistent), freelancer has mu = ln(30) with sigma = 0.6 (faster but more variable). These priors are used '
    'by fitDistribution() when no historical payment data exists for the debtor.'
)

# ── 4.8 Seasonal Patterns ────────────────────────────────────────────────────

doc.add_heading('4.8 Seasonal Patterns', level=2)

doc.add_paragraph(
    'Seasonal payment pattern integration (Gap 13) adjusts the P(Pay) distribution for known monthly variations. '
    'Patterns come from two sources. Riley extracts seasonal intelligence from user conversations (e.g., "construction '
    'companies slow down in December") and stores them as aiFacts with category=\'seasonal_pattern\'. The system also '
    'learns patterns from payment history via calculateLearnedSeasonalPatterns(), which requires 12+ months of data '
    'and 2+ data points per month, identifying months that deviate more than 20% from the overall average.'
)

doc.add_paragraph(
    'Three adjustment types shift the distribution\'s mu parameter: slow months (+0.15, approximately 15% later payment), '
    'fast months (\u22120.15, 15% earlier), and year-end months (\u22120.20, the strongest pull-forward reflecting year-end '
    'payment runs). getEffectiveSeasonalAdjustments() merges Riley-captured and learned patterns, with learned patterns '
    'overriding Riley when learning confidence exceeds 0.7. All four callers of fitDistribution() \u2014 the adaptive '
    'scheduler, action planner, weekly review service, and Riley assistant \u2014 pass seasonal adjustments through.'
)

# ── 4.9 Debtor Grouping ──────────────────────────────────────────────────────

doc.add_heading('4.9 Debtor Grouping', level=2)

doc.add_paragraph(
    'Debtor grouping (Gap 12) handles the real-world situation where multiple contacts at the same company share an '
    'accounts payable department. Sending a friendly email to one contact and a firm email to another on the same day '
    'is contradictory and undermines credibility. The debtorGroups table links contacts, and the action planner\'s '
    'enforceDebtorGroupConsistency() function runs as a post-planning sweep.'
)

doc.add_paragraph(
    'The sweep queries all same-day scheduled actions for contacts in the same group. It aligns all actions to the '
    'highest tone level in the group (tone only escalates, never de-escalates within a group day). It then cancels '
    'duplicate actions, keeping the highest-priority action and marking others with cancellationReason = '
    '\'debtor_group_same_day_conflict\'. Riley\'s detectPotentialGroups() function suggests groups automatically by '
    'matching email domains (excluding 14 known generic providers like gmail.com, outlook.com, etc.).'
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 5. MESSAGE GENERATION PIPELINE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('5. Message Generation Pipeline', level=1)

doc.add_paragraph(
    'Charlie does not use templates for normal operations. Every email, SMS, and voice script is generated fresh by '
    'an LLM (Claude, via server/services/llm/claude.ts) with full business context. Templates exist only as a fallback '
    'when the LLM is unavailable. This section describes the generation pipeline end to end.'
)

# ── 5.1 LLM Generation ───────────────────────────────────────────────────────

doc.add_heading('5.1 LLM Generation', level=2)

doc.add_paragraph(
    'The AI message generator (server/services/aiMessageGenerator.ts) exposes three methods: generateEmail(), '
    'generateSMS(), and generateVoiceScript(). Each takes a MessageContext (customer name, invoice details, '
    'outstanding amounts, prior contact history, promise-to-pay status, VIP flags, dispute status, tenant details, '
    'and payment link) and ToneSettings (playbook stage, tone profile, reason code, template ID, tenant style, '
    'and late payment legislation flag).'
)

doc.add_paragraph(
    'Each method follows the same flow. First, canAttemptGeneration() queries the circuit breaker. If the circuit '
    'is closed (healthy), the generator builds a system prompt and user prompt specific to the channel, then calls '
    'callLLMWithRetry() which makes two attempts with a 10-second delay between failures. The LLM is asked to return '
    'structured JSON (subject, body, callToAction for email; body for SMS; voiceScript for voice). If generation '
    'succeeds, the output passes through validation. If both attempts fail, recordFailure() updates the circuit breaker.'
)

doc.add_paragraph(
    'System prompts encode the tone rules. For email, the prompt specifies the sender persona (agent name and title '
    'from the tenant\'s agent persona configuration), the tone level with behavioural guidelines, and strict formatting '
    'rules (HTML paragraph tags, no signatures for SMS, no mentioning AI identity). The user prompt provides all '
    'invoice details, payment history, and context signals. For consolidated invoices with multiple line items, '
    'the email system prompt instructs the LLM to include an HTML table.'
)

# ── 5.2 Output Validation ────────────────────────────────────────────────────

doc.add_heading('5.2 Output Validation', level=2)

doc.add_paragraph(
    'Every generated message passes through validateGeneratedMessage() in server/services/llmOutputValidator.ts, '
    'which applies five checks.'
)

add_bold_para('Rule 1: Length bounds. ', 'Email content must be 200\u20135,000 characters (after HTML tag stripping). '
    'SMS must be 50\u2013160 characters. Voice scripts must be 100\u20132,000 characters. Messages outside these bounds '
    'are rejected.')

add_bold_para('Rule 2: Debtor name presence. ', 'The message must contain either the debtor\'s full name or at least '
    'the first word of their name (minimum 2 characters). This check is case-insensitive and skipped if the debtor '
    'name is the generic "Customer".')

add_bold_para('Rule 3: Invoice reference or currency amount. ', 'The message must contain at least one invoice reference '
    '(case-insensitive match) or a currency amount matching the pattern \u00a3[digits]. This ensures every collections '
    'message is specific enough for the debtor to identify what they owe.')

add_bold_para('Rule 4: System prompt leakage. ', 'Eight regex patterns detect AI identity exposure: "as an ai", '
    '"language model", "i\'m claude", "i am claude", "anthropic", "system prompt", and format markers from other '
    'LLMs ([INST], <|im_start|>). Any match fails validation immediately.')

add_bold_para('Rule 5: Tone alignment. ', 'Friendly-tone messages are rejected if they contain "legal proceedings", '
    '"consequences", or "failure to". Professional-tone messages are rejected for "legal proceedings". Legal-tone '
    'messages must contain at least one of "pre-action", "proceedings", "statutory interest", or "civil procedure". '
    'Firm and formal tones generate warnings but do not fail validation.')

doc.add_paragraph(
    'If validation fails, the generator retries the LLM once. If the retry also fails validation, the system falls '
    'back to template content (for email and SMS) or returns an error (for voice, which has no template fallback).'
)

# ── 5.3 Circuit Breaker ──────────────────────────────────────────────────────

doc.add_heading('5.3 Circuit Breaker', level=2)

doc.add_paragraph(
    'The LLM circuit breaker (server/services/llmCircuitBreaker.ts) implements a three-state machine per tenant to '
    'protect the system against sustained LLM outages. The states are closed (healthy), open (LLM unavailable), '
    'and half_open (recovery probe in progress).'
)

doc.add_paragraph(
    'In the closed state, all generation attempts are allowed. When FAILURE_THRESHOLD (3) failures occur within '
    'FAILURE_WINDOW_MS (30 minutes), the circuit transitions to open. In the open state, generation is blocked. '
    'Every PROBE_INTERVAL_MS (5 minutes), a single probe request is allowed, transitioning the circuit to half_open. '
    'If the probe succeeds, the circuit returns to closed and admin notifications are sent. If the probe fails, '
    'the circuit returns to open.'
)

doc.add_paragraph(
    'After TEMPLATE_ACTIVATION_MS (4 hours) in the open state, canAttemptGeneration() returns useTemplate = true, '
    'signalling the message generator to use static template fallbacks rather than queuing actions indefinitely. '
    'Admin notifications are sent via email when the circuit opens or closes. SMS notifications are defined but '
    'non-functional (the users table lacks a phone column).'
)

# ── 5.4 Template Fallback ────────────────────────────────────────────────────

doc.add_heading('5.4 Template Fallback', level=2)

doc.add_paragraph(
    'The template fallback library (server/services/templateFallback.ts) provides 10 static templates: 5 tones '
    '(friendly, professional, firm, formal, legal) across 2 channels (email and SMS). Voice has no template fallback '
    '\u2014 voice scripts require the nuance and context-awareness that only LLM generation can provide.'
)

doc.add_paragraph(
    'Templates use {{variable}} placeholders substituted by the substitute() function. Variables include contactName, '
    'companyName, totalOutstanding, oldestInvoiceRef, oldestInvoiceDays, creditorName, agentName, agentTitle, '
    'agentEmail, and agentPhone. The buildTemplateContext() function assembles these from the message context and '
    'the tenant\'s active agent persona (fetched via storage.getActiveAgentPersona()). SMS templates are truncated '
    'to 160 characters and use first names only; tenant names are abbreviated (removing Ltd, PLC, Inc suffixes). '
    'Email templates use HTML paragraph tags for formatting.'
)

doc.add_paragraph(
    'The legal-tone email template is marked as "LETTER BEFORE ACTION" \u2014 the formal pre-action notice required '
    'under UK debt collection practice. All template-generated messages are flagged with generationMethod = '
    '\'template_fallback\' so they can be distinguished from LLM-generated content in analytics and audit.'
)

# ── 5.5 Compliance Engine ────────────────────────────────────────────────────

doc.add_heading('5.5 Compliance Engine', level=2)

doc.add_paragraph(
    'The compliance engine (server/services/compliance/complianceEngine.ts) is the final gate before any message '
    'reaches the delivery pipeline. checkCompliance() applies six rules and returns a ComplianceResult with an '
    'action directive: send, block, regenerate, or queue_for_approval.'
)

add_bold_para('Rule 1: Frequency cap. ', 'Counts outbound actions (email, SMS, call with status completed/sent/delivered, '
    'excluding failed and bounced deliveries) within the contact window (tenant setting, default 14 days). '
    'If the count meets or exceeds maxTouchesPerWindow (default 3), the action is queued for approval.')

add_bold_para('Rule 2: Channel cooldown. ', 'Checks whether an email was sent to this contact within the email '
    'cooldown period (tenant setting, default 3 days). If so, the action is queued for approval.')

add_bold_para('Rule 3: Time-of-day. ', 'Checks whether the current time falls within the tenant\'s configured '
    'business hours (default 08:00\u201318:00). If outside business hours, the action is queued for approval.')

add_bold_para('Rule 4: Prohibited language. ', 'Scans the email subject and body against three pattern groups. '
    'Legal threat patterns (court action, legal proceedings, solicitor, CCJ, statutory demand, winding up, bailiff, '
    'enforcement agent) are only allowed at formal tone level \u2014 at any other tone, they trigger a block. '
    'Harassment patterns (pay immediately or, failure to pay will result, no choice but, final warning, last chance) '
    'are always blocked. Profanity patterns are always blocked.')

add_bold_para('Rule 5: Data isolation. ', 'checkDataIsolation() queries all invoices for the tenant and checks '
    'whether the email content references any invoice belonging to a different contact. Invoice numbers shorter than '
    '3 characters are skipped to avoid false positives. Word-boundary regex matching prevents substring matches. '
    'A violation triggers an immediate block \u2014 cross-debtor data leakage is a multi-tenancy breach.')

add_bold_para('Rule 6: Debtor vulnerability. ', 'If the contact is flagged as isPotentiallyVulnerable and the tone '
    'level exceeds professional, the action is sent for regeneration at a lower tone. This aligns with the tone '
    'escalation engine\'s VULNERABLE_CEILING.')

doc.add_paragraph(
    'Action determination follows a priority order. If prohibited language or data isolation is violated, the action '
    'is blocked. If vulnerability is violated, the message is regenerated. If only time-of-day is violated, the action '
    'is queued for approval (can be sent during next business hours). Frequency cap and cooldown violations also queue '
    'for approval. Every compliance check is logged to the complianceChecks table with the result, rules checked, '
    'violations, and agent reasoning \u2014 creating a full audit trail.'
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 6. EXECUTION & DELIVERY PIPELINE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('6. Execution & Delivery Pipeline', level=1)

# ── 6.1 Scheduler ────────────────────────────────────────────────────────────

doc.add_heading('6.1 Scheduler Architecture', level=2)

doc.add_paragraph(
    'The collections scheduler (server/services/collectionsScheduler.ts) implements a two-phase architecture. '
    'Phase 1 (the planner) runs every 60 minutes, calling actionPlanner.planActionsForAllTenants() to scan all enabled '
    'tenants, evaluate overdue invoices, and create scheduled actions. Phase 2 (the executor) runs every 10 minutes, '
    'calling actionExecutor.executeScheduledActions() to process up to 50 actions per run that have reached their '
    'scheduled time and have been approved.'
)

doc.add_paragraph(
    'The portfolio controller (server/services/portfolioController.ts) adds a higher-level orchestration layer using '
    'cron scheduling. The nightly urgency task runs at 2 AM daily, calling runNightly() to recompute urgency and '
    'per-debtor weights. The adaptive planning task runs every 6 hours, calling planAdaptiveActions() for each active '
    'workflow. The daily plan generation task also runs at 2 AM (Europe/London timezone), creating a daily plan for '
    'each tenant with collections automation enabled. Both the scheduler and portfolio controller start from the '
    'server\'s bootstrap sequence in server/startup/orchestrator.ts.'
)

# ── 6.2 Validation Gate ──────────────────────────────────────────────────────

doc.add_heading('6.2 Execution-Time Validation Gate', level=2)

doc.add_paragraph(
    'The validation gate in actionExecutor.ts is critical. Actions may be planned hours or days before execution, and '
    'the world changes in the interim. validateActionBeforeExecution() re-checks every assumption before spending '
    'resources on message generation and delivery.'
)

doc.add_paragraph(
    'The gate performs five checks in order. First, it checks the legal response window (Gap 10): if the contact has '
    'a legalResponseWindowEnd date that is still in the future, the action is cancelled with reason '
    '\'legal_response_window_active\'. If the window has expired but hasn\'t been explicitly resolved by a manager, '
    'the action is cancelled with \'legal_response_window_expired_pending_resolution\'. Second, it checks for probable '
    'payments (Gap 14): if a bank transaction has been matched to this debtor with high or medium confidence, the action '
    'is cancelled with \'probable_payment_detected\'. Third, it re-reads all invoices in the action\'s bundle from the '
    'database, filtering out any that have become paid, void, deleted, disputed, on hold, or paused since planning. '
    'If all invoices are excluded, the action is cancelled. If the bundle has changed (fewer valid invoices), it is '
    'cancelled with \'bundle_modified_requires_replan\'. Fourth, if the remaining total falls below the tenant\'s '
    'minimum chase threshold, the action is cancelled.'
)

# ── 6.3 Delivery Tracking ────────────────────────────────────────────────────

doc.add_heading('6.3 Delivery Tracking', level=2)

doc.add_paragraph(
    'Gap 8 implemented end-to-end delivery tracking for SendGrid emails. When the action executor sends an email via '
    'the SendGrid wrapper, it attaches custom arguments (tenant_id, action_id, contact_id, invoice_id) to the API '
    'call. When SendGrid fires delivery webhooks (delivered, bounced, dropped, opened, clicked), the webhook handler '
    'extracts these custom args and publishes events through the event bus. processDeliveryOutcome() updates the '
    'action\'s delivery status. processEffectivenessUpdate() in the channel effectiveness service recalculates the '
    'debtor\'s channel scores.'
)

doc.add_paragraph(
    'Two delivery tracking gaps remain. Connected email (Gmail/Outlook OAuth path) has no equivalent to SendGrid '
    'custom args or webhooks \u2014 delivery tracking is SendGrid-only. Several SendGrid event types (processed, '
    'unsubscribe, spamreport, group_unsubscribe, group_resubscribe) are recorded in contactOutcomes but not '
    'processed by the delivery event pipeline.'
)

# ── 6.4 Retry Logic ──────────────────────────────────────────────────────────

doc.add_heading('6.4 Retry Logic', level=2)

doc.add_paragraph(
    'When email delivery fails, handleSendFailure() in the action executor implements a two-retry strategy with '
    'increasing delays. The first retry is scheduled 5 minutes after the failure; the second retry 30 minutes after. '
    'Each retry creates a new action record with retryOf pointing to the original action ID and an incremented '
    'retryCount. The original action is marked as status=\'failed\' with deliveryStatus=\'failed\'.'
)

doc.add_paragraph(
    'If both retries fail, the action is marked status=\'failed_permanent\' with deliveryStatus=\'failed_permanent\'. '
    'A hard bounce event is published through the event bus, which triggers a timeline event for Data Health visibility '
    'and updates the channel effectiveness score. The maximum retry count of 2 prevents the system from hammering '
    'a permanently invalid email address.'
)

# ── 6.5 Legal Evidence & Voice ────────────────────────────────────────────────

doc.add_heading('6.5 Legal Evidence & Voice Contact Records', level=2)

doc.add_paragraph(
    'Every action in the actions table includes the complete context used to generate it: the message content, '
    'the tone level (agentToneLevel), the scheduling reasoning, and the invoices covered. This creates a legally '
    'defensible audit trail showing exactly what was communicated, when, why, and to whom. The compliance checks '
    'table adds a parallel audit trail of every rule that was evaluated.'
)

doc.add_paragraph(
    'For voice calls, the action executor calls sendVoiceCall() in server/services/communications/sendVoiceCall.ts, '
    'which wraps the Retell AI API. The voiceContactRecord (JSONB column on the actions table) stores the call outcome '
    'metadata. Voice calls require an agentId (Retell AI agent configuration) and are subject to the same communication '
    'mode enforcement as email and SMS.'
)

doc.add_paragraph(
    'The legal response window (Gap 10) is set by setLegalResponseWindowIfNeeded() after any legal-tone action '
    'completes. It writes legalResponseWindowEnd = now + 30 days on the contact record and creates a timeline event. '
    'The daily legalWindowJob.ts monitors these windows: at day 25, it creates an expiry warning event (deduped '
    'against the previous 7 days). At day 30+, it creates an expired event. Both events are flagged as '
    'outcomeRequiresReview = true. Three resolution paths are available via API: resume_collections (clear the window), '
    'refer_debt_recovery (keep the window, log handoff), or extend_window (new 30-day window).'
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 7. COMPLIANCE & SAFETY SYSTEMS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('7. Compliance & Safety Systems', level=1)

doc.add_paragraph(
    'Charlie operates in a heavily regulated domain. UK debt collection is governed by the Financial Conduct Authority, '
    'the "Time to Pay Up" legislation (effective 24 March 2026), and industry codes of practice. Every safety mechanism '
    'in Charlie exists because a specific legal or ethical risk demanded it.'
)

doc.add_heading('Communication Mode Enforcement', level=2)

doc.add_paragraph(
    'All outbound communications pass through central enforcement wrappers that check the tenant\'s communication mode '
    'before anything reaches the wire. Four modes exist: Off (all sends blocked), Testing (sends redirected to test '
    'addresses with [TEST] prefix), Soft Live (currently identical to Testing until contact-level opt-in is built), '
    'and Live (full production delivery). The wrappers are in server/services/sendgrid.ts (enforceCommunicationMode), '
    'server/services/vonage.ts, and server/services/communications/sendVoiceCall.ts. Critically, all wrappers fail '
    'closed: if the database is unreachable during the mode check, sends are blocked, not allowed through. New tenants '
    'default to Testing mode.'
)

doc.add_heading('30-Day Statutory Response Window', level=2)

doc.add_paragraph(
    'After any legal-tone action is delivered, a 30-day statutory response window is automatically set on the contact. '
    'During this window, the execution-time validation gate blocks all further collection actions for that debtor. '
    'When the window expires without resolution, it continues to block actions \u2014 a manager must explicitly resolve '
    'the window via the API (choosing to resume collections, refer to debt recovery, or extend for another 30 days). '
    'This prevents the system from inadvertently sending informal follow-ups after a formal pre-action notice, which '
    'would undermine the legal position.'
)

doc.add_heading('Channel Preference Overrides', level=2)

doc.add_paragraph(
    'Debtor channel preferences are hard overrides that cannot be bypassed by algorithmic channel selection. '
    'If a debtor says "don\'t call me", voiceEnabled is set to false in customerPreferences, and Charlie will never '
    'select voice for that debtor regardless of what the channel selection algorithm recommends. Preferences can be '
    'set manually by users, or extracted automatically by Riley from conversation analysis. The system respects null '
    'as enabled (backwards compatible) and falls back through email, SMS, and voice in order when the preferred '
    'channel is disabled.'
)

doc.add_heading('Tone Velocity Cap', level=2)

doc.add_paragraph(
    'The tone velocity cap prevents aggressive escalation by limiting tone changes to plus or minus one level per '
    'contact cycle. A debtor cannot go from friendly to firm in a single step \u2014 they must pass through '
    'professional first. This creates a graduated, defensible escalation path. The significant payment override '
    '(50%+ of outstanding paid) is the only exception, and it only works downward (resetting to professional).'
)

doc.add_heading('Cadence Limits', level=2)

doc.add_paragraph(
    'The compliance engine enforces a maximum of 3 touches per 14-day window (both configurable per tenant) and '
    'a 3-day cooldown between emails to the same contact. Failed deliveries (bounced, dropped) are excluded from '
    'the touch count \u2014 they should not consume the debtor\'s contact allowance. Time-of-day enforcement ensures '
    'contacts are only made during business hours (default 08:00\u201318:00).'
)

doc.add_heading('Minimum Chase Threshold', level=2)

doc.add_paragraph(
    'If the consolidated outstanding amount for a debtor falls below the tenant\'s minimum chase threshold, '
    'the action is skipped. This prevents sending collection communications for trivially small amounts, which '
    'damages the creditor\'s reputation more than the amount is worth recovering.'
)

doc.add_heading('Vulnerable Debtor Protection', level=2)

doc.add_paragraph(
    'Debtors flagged as potentially vulnerable (via the isPotentiallyVulnerable field on contacts) receive two '
    'layers of protection. The tone escalation engine caps their tone at professional (VULNERABLE_CEILING), and '
    'the compliance engine rejects any message at firm or formal tone for vulnerable debtors, sending it for '
    'regeneration at a lower tone. These protections cannot be overridden by algorithmic decisions.'
)

doc.add_heading('Circuit Breaker & Admin Alerts', level=2)

doc.add_paragraph(
    'When the LLM circuit breaker opens (3 failures within 30 minutes), admin users receive email notifications. '
    'The system queues actions for the first 4 hours, then falls back to templates. When the circuit closes '
    '(successful probe), admins are notified again. This ensures human oversight of AI capability degradation '
    'without requiring manual monitoring.'
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 8. DATA ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('8. Data Architecture', level=1)

doc.add_paragraph(
    'Charlie\'s data model spans approximately 15 key tables in a PostgreSQL database (Neon serverless), managed '
    'by Drizzle ORM with the full schema defined in shared/schema.ts. Every table includes a tenantId foreign key '
    'for multi-tenant data isolation.'
)

add_table(
    ['Table', 'Purpose', 'Key Columns'],
    [
        ['actions', 'Every collection action (planned, executed, completed, failed)',
         'id, tenantId, contactId, invoiceId, actionType, status, scheduledFor, agentToneLevel, deliveryStatus, retryCount, retryOf, voiceContactRecord, cancellationReason'],
        ['customerLearningProfiles', 'Per-debtor learning state',
         'tenantId, contactId, emailEffectiveness, smsEffectiveness, voiceEffectiveness, learningConfidence, preferredChannel, totalInteractions, prsRaw, prsConfidence, debtorUrgency, contributionWeight, trendMultiplier, urgencyUpdatedAt'],
        ['customerBehaviorSignals', 'Historical payment behaviour metrics',
         'tenantId, contactId, medianDaysToPay, p75DaysToPay, volatility, trend, emailReplyRate, smsReplyRate, weekdayEffect, promiseBreachCount'],
        ['debtorIntelligence', 'Companies House enrichment + credit risk',
         'tenantId, contactId (unique), companyNumber, companyStatus, sicCodes, companyAge, companySize, filingHealth, lateFilings, insolvencyRisk, creditRiskScore, aiRiskSummary, enrichedAt, enrichmentSources'],
        ['debtorGroups', 'Linked debtor entities (shared AP department)',
         'id, tenantId, name, members (contactIds), createdAt'],
        ['aiFacts', 'Business intelligence from Riley conversations',
         'tenantId, entityId, category, factKey, factValue, confidence, source, sourceConversationId, expiresAt, isActive'],
        ['contactOutcomes', 'Raw delivery events from webhooks',
         'tenantId, contactId, actionId, eventType, channel, metadata, occurredAt'],
        ['paymentPromises', 'PTP records for reliability tracking',
         'tenantId, contactId, invoiceId, promiseType, promisedDate, promisedAmount, status, evaluatedAt, sourceType, channel'],
        ['contacts', 'Debtor/customer records with AR overlay',
         'tenantId, name, email, phone, arContactEmail, arContactPhone, arContactName, arNotes, riskScore, isPotentiallyVulnerable, legalResponseWindowEnd, debtorGroupId, creditLimit'],
        ['customerPreferences', 'Channel opt-outs and contact windows',
         'tenantId, contactId, emailEnabled, smsEnabled, voiceEnabled, channelPreferenceSource, channelPreferenceNotes, tradingName'],
        ['invoices', 'Main invoice data',
         'tenantId, contactId, invoiceNumber, amount, dueDate, status, reminderCount, pauseState, xeroInvoiceId'],
        ['complianceChecks', 'Audit log of every compliance evaluation',
         'tenantId, contactId, actionId, checkResult, rulesChecked, violations, agentReasoning, reviewedBy, reviewedAt'],
        ['timelineEvents', 'Chronological record of all debtor interactions',
         'tenantId, customerId, occurredAt, direction, channel, summary, preview, createdByType, outcomeRequiresReview'],
        ['agentPersonas', 'AI persona configuration per tenant',
         'tenantId, name, title, email, phone, style, isActive'],
        ['probablePayments', 'Unreconciled bank transaction matches',
         'tenantId, contactId, invoiceId, transactionId, confidence, amount, matchedAt'],
    ]
)

doc.add_paragraph(
    'The data flow between these tables forms Charlie\'s learning loop. Invoices and contacts flow in from Xero sync. '
    'The decision engine reads invoices, contacts, customerBehaviorSignals, and customerLearningProfiles to make '
    'decisions. Actions are written with full context. Delivery webhooks update contactOutcomes. Payment detection '
    'triggers attribution analysis that updates customerLearningProfiles. The nightly job recalculates PRS, urgency, '
    'and enrichment scores. Tomorrow, the decision engine reads the updated profiles and makes better decisions.'
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 9. INTEGRATION POINTS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('9. Integration Points', level=1)

doc.add_heading('Xero (Accounting Platform)', level=2)

doc.add_paragraph(
    'Xero is Charlie\'s primary data source for invoices and contacts. The integration (server/services/xeroSync.ts) '
    'uses OAuth 2.0 with offline_access scope for token refresh. Sync runs on a 4-hour background cadence with two modes: '
    'INITIAL (clean sweep for first connection) and ONGOING (upsert, never delete, never overwrite AR overlay fields). '
    'Rate limiting enforces 1.5-second delays between paginated API calls and 60-second waits on 429 responses. '
    'Token refresh is proactive (2-minute buffer check before API calls) with reactive 401 retry fallback. '
    'Failed refresh marks the connection as "expired" and prompts the user to reconnect. Invoice PAID transitions '
    'detected during sync trigger payment attribution through the event bus. Bank transaction sync methods exist in '
    'the API client but are not yet wired into the sync pipeline (Gap 14).'
)

doc.add_heading('SendGrid (Email Delivery)', level=2)

doc.add_paragraph(
    'SendGrid handles all transactional email delivery. Outbound emails pass through enforceCommunicationMode() '
    'in server/services/sendgrid.ts, which checks the tenant\'s communication mode before calling the SendGrid API. '
    'Custom arguments (tenant_id, action_id, contact_id, invoice_id) are attached to every outbound email for '
    'webhook correlation. Inbound email parsing via SendGrid webhooks (/api/webhooks/sendgrid/inbound) matches '
    'replies to debtors via reply tokens or email addresses, triggers intent extraction via Claude, and creates '
    'behavioural signals. Delivery webhooks flow through the event bus to update action delivery status and '
    'channel effectiveness scores.'
)

doc.add_heading('Vonage (SMS)', level=2)

doc.add_paragraph(
    'Vonage provides SMS delivery with communication mode enforcement in server/services/vonage.ts. SMS messages '
    'are constrained to 160 characters. The integration supports both standard SMS and WhatsApp messaging channels.'
)

doc.add_heading('Retell AI (Voice)', level=2)

doc.add_paragraph(
    'Retell AI powers voice calls through server/services/communications/sendVoiceCall.ts. Voice calls require '
    'an agent configuration (agentId) in Retell and pass dynamic variables (customer name, invoice details) to the '
    'voice AI. Communication mode enforcement is applied. Voice scripts are LLM-generated with no template fallback '
    '\u2014 the nuance required for a natural conversation cannot be captured in static text.'
)

doc.add_heading('Claude API (LLM)', level=2)

doc.add_paragraph(
    'All AI generation uses Anthropic\'s Claude API via server/services/llm/claude.ts, which provides generateText() '
    'and generateJSON() abstractions. Claude is used for: collections email/SMS/voice generation (aiMessageGenerator.ts), '
    'intent extraction from debtor replies (intentAnalyst.ts), Riley assistant conversations (rileyAssistant.ts), '
    'AI risk summaries in debtor enrichment (debtorEnrichmentService.ts), and action plan optimisation (actionPlanner.ts). '
    'The circuit breaker protects against API outages. All Claude usage is inference on customer-owned data \u2014 '
    'Xero API data is never used for model training per Xero\'s March 2026 developer terms.'
)

doc.add_heading('Companies House (Enrichment)', level=2)

doc.add_paragraph(
    'The Companies House API (server/services/companiesHouse.ts) provides company registration data, filing history, '
    'and incorporation details for UK companies. Authentication uses HTTP Basic with the API key as username. '
    'Three endpoints are used: company search (/search/companies), company profile (/company/{number}), and filing '
    'history (/company/{number}/filing-history). Rate limit is 600 requests per 5 minutes. The API is free '
    '(register at developer.company-information.service.gov.uk). The service degrades gracefully if the API key '
    'is not configured.'
)

doc.add_heading('Open Banking (Future)', level=2)

doc.add_paragraph(
    'Open Banking integration via TrueLayer or Yapily is planned for Gap 14 (unreconciled payment detection). '
    'Bank transaction data has no AI/ML training restrictions (unlike Xero data) and will be the primary source for '
    'model training. The two-source architecture is a hard architectural boundary: Open Banking data trains models, '
    'Xero API data is used for operational inference only.'
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 10. WHAT MAKES THIS SYSTEM EXCEPTIONAL
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('10. What Makes This System Exceptional', level=1)

doc.add_heading('Closed-Loop Learning', level=2)

doc.add_paragraph(
    'Most collections systems are open-loop: they send messages based on rules and never learn whether those messages '
    'worked. Charlie closes the loop. Every delivery event, every engagement signal, every payment outcome feeds back '
    'into the debtor\'s learning profile. Channel effectiveness scores update via adaptive EMA. Payment attribution '
    'credits the right channels. PRS tracks promise reliability. The system that runs today is measurably different '
    'from the system that ran last month \u2014 and measurably better.'
)

doc.add_heading('Statistical Modelling', level=2)

doc.add_paragraph(
    'Charlie does not guess when debtors will pay. It fits a log-normal distribution to each debtor\'s historical '
    'payment data, incorporating trend, volatility, p75 spread, and seasonal adjustments. The conditional CDF '
    'gives a mathematically grounded probability of payment within any time horizon. Three-scenario forecasting '
    '(25th/50th/75th percentiles) feeds directly into cashflow projections. This is not a heuristic \u2014 it is '
    'a proper statistical model that improves as data accumulates.'
)

doc.add_heading('Bayesian Reasoning', level=2)

doc.add_paragraph(
    'The Promise Reliability Score uses Bayesian inference to handle the thin-data problem that plagues every '
    'scoring system. A debtor who has kept one promise out of one does not get a perfect score \u2014 the Bayesian '
    'prior (k=3, prior=60) regresses them toward the population mean. As evidence accumulates, the prior\'s influence '
    'diminishes and the score converges on reality. Recency weighting ensures that a debtor\'s score today reflects '
    'their behaviour today, not their behaviour two years ago.'
)

doc.add_heading('Graduated Autonomy', level=2)

doc.add_paragraph(
    'Charlie operates on a spectrum from full human control to full autonomy. Communication modes (Off, Testing, '
    'Soft Live, Live) control whether messages actually reach debtors. Semi-auto mode requires human approval for '
    'every action. Full-auto mode lets Charlie execute independently. Human review flags trigger on low confidence, '
    'high-value first contacts, and VIP customers. The minimum chase threshold, velocity cap, and vulnerable debtor '
    'ceiling all constrain autonomy. The system is designed to earn trust incrementally.'
)

doc.add_heading('Legal Defensibility', level=2)

doc.add_paragraph(
    'Every action Charlie takes is auditable. The actions table records the full context, tone level, and reasoning. '
    'The complianceChecks table records every rule evaluation. Timeline events create a chronological narrative. '
    'The 30-day statutory response window enforces the prescribed response period after legal communications. '
    'Prohibited language scanning prevents threats at inappropriate tone levels. Data isolation checks prevent '
    'cross-debtor information leakage. If challenged in court, the creditor can demonstrate that every communication '
    'followed a defensible, documented process.'
)

doc.add_heading('Two-Source Architecture', level=2)

doc.add_paragraph(
    'Charlie enforces a hard boundary between data sources. Xero API data (invoices, contacts, bank transactions) '
    'is used for operational inference \u2014 matching payments, calculating overdue amounts, enriching debtor profiles. '
    'It is never used for AI/ML model training, in compliance with Xero\'s March 2026 developer terms. Open Banking '
    'data (via TrueLayer/Yapily) has no such restriction and will be the exclusive source for model training. '
    'This architectural separation is not just a policy \u2014 it is enforced at the code level.'
)

doc.add_heading('Resilience', level=2)

doc.add_paragraph(
    'Charlie is designed to degrade gracefully, not fail catastrophically. If the LLM is unavailable, the circuit '
    'breaker activates and the system falls back to templates after 4 hours. If Xero sync fails, existing data '
    'continues to drive decisions. If Companies House is unreachable, enrichment is skipped and credit risk scoring '
    'uses available data. If the database is unreachable during communication mode checks, all sends are blocked '
    '(fail closed, not fail open). Communication wrappers, validation gates, and compliance checks all use non-fatal '
    'error handling \u2014 a failure in one subsystem does not cascade to crash the server. The unhandledRejection '
    'handler logs but does not exit, ensuring background task failures are visible but not destructive.'
)

doc.add_paragraph(
    'Charlie is not a finished product. The architecture described in this document represents version 1.1, with '
    '14 engineering gaps identified and implemented during the Charlie hardening sprint. The next phase brings '
    'Open Banking integration (Gap 14), connected email delivery tracking, Data Health bounce integration, and '
    'the evolution from a single decision engine to a five-agent architecture with shared state and inter-agent '
    'orchestration. The foundation is solid. The ceiling is high.'
)


# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════

output_path = os.path.expanduser('~/Documents/qashivo/CHARLIE_TECHNICAL_ARCHITECTURE.docx')
doc.save(output_path)
print(f'Document saved to {output_path}')
print(f'File size: {os.path.getsize(output_path):,} bytes')
